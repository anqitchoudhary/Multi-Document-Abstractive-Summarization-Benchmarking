#                                                                               Code for CNN/Daily Mail Dataset

# ==============================================================================
#
# COMPREHENSIVE BENCHMARK FOR ABSTRACTIVE SUMMARIZATION
#
# This script provides a full pipeline for fine-tuning and evaluating ten
# different transformer-based models on the CNN/Daily Mail dataset.
#
#
# ==============================================================================

# ==============================================================================
# SECTION 1: SETUP & IMPORTS
# ==============================================================================
import os
import torch
import numpy as np
import pandas as pd
import spacy
from datasets import load_dataset
from transformers import (
    AutoTokenizer,
    AutoModelForSeq2SeqLM,
    Seq2SeqTrainingArguments,
    Seq2SeqTrainer,
    DataCollatorForSeq2Seq,
    LEDForConditionalGeneration,
    LongT5ForConditionalGeneration,
)
import evaluate
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.neighbors import NearestNeighbors
import argparse
import logging
import time

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Set random seed for reproducibility
torch.manual_seed(42)
np.random.seed(42)

# ==============================================================================
# SECTION 2: GLOBAL CONFIGURATION & MODEL DEFINITIONS
# ==============================================================================

# --- Hardware Configuration ---
DEVICE = "cuda:0" if torch.cuda.is_available() else "cpu"
logging.info(f"Using device: {DEVICE}")
if torch.cuda.is_available():
    logging.info(f"GPU Model: {torch.cuda.get_device_name(DEVICE)}")

# --- Dataset Configuration ---
DATASET_NAME = "ccdv/cnn_dailymail"
DATASET_VERSION = "3.0.0"

# --- Model Configurations ---
# This dictionary drives the entire benchmark process.
# Each entry defines a model, its checkpoint, specific parameters, and type.
MODEL_CONFIGS = {
    "bart": {
        "model_class": AutoModelForSeq2SeqLM,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "facebook/bart-large-cnn",
        "type": "sds", # Single-Document Summarization
        "prefix": "",
        "max_input_length": 1024,
    },
    "pegasus": {
        "model_class": AutoModelForSeq2SeqLM,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "google/pegasus-cnn_dailymail",
        "type": "sds",
        "prefix": "",
        "max_input_length": 1024,
    },
    "t5-base": {
        "model_class": AutoModelForSeq2SeqLM,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "t5-base",
        "type": "sds",
        "prefix": "summarize: ",
        "max_input_length": 512,
    },
    "t5-large": {
        "model_class": AutoModelForSeq2SeqLM,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "t5-large",
        "type": "sds",
        "prefix": "summarize: ",
        "max_input_length": 512,
    },
    "led": {
        "model_class": LEDForConditionalGeneration,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "allenai/led-large-16384",
        "type": "sds",
        "prefix": "",
        "max_input_length": 4096, # Leverage long-context capability
    },
    "long-t5": {
        "model_class": LongT5ForConditionalGeneration,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "google/long-t5-tglobal-base",
        "type": "sds",
        "prefix": "summarize: ",
        "max_input_length": 4096, # Leverage long-context capability
    },
    "primera": {
        "model_class": LEDForConditionalGeneration, # PRIMERA is based on LED
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "allenai/PRIMERA-arxiv",
        "type": "mds", # Multi-Document Summarization
        "prefix": "",
        "max_input_length": 8192, # For concatenated documents
    },
    "bart-entity": {
        "model_class": AutoModelForSeq2SeqLM,
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "facebook/bart-large-cnn",
        "type": "sds_entity", # Special type for knowledge-enhancement
        "prefix": "",
        "max_input_length": 1024,
    },
    "tg-multisum": {
        "model_class": "HGSUM", # Custom implementation required
        "tokenizer_class": AutoTokenizer,
        "checkpoint": "allenai/PRIMERA-arxiv", # HGSUM uses PRIMERA as its base
        "type": "mds_custom", # Custom Multi-Document Summarization
        "prefix": "",
        "max_input_length": 8192,
    },
    "dca": {
        "model_class": "DCA", # Custom implementation required
        "tokenizer_class": AutoTokenizer,
        "checkpoint": None, # Trained from scratch
        "type": "mds_custom",
        "prefix": "",
        "max_input_length": 8192,
    },
    "absformer": {
        "model_class": "Absformer", # Custom implementation required
        "tokenizer_class": AutoTokenizer,
        "checkpoint": None, # Unsupervised, trained from scratch
        "type": "mds_custom_unsupervised",
        "prefix": "",
        "max_input_length": 8192,
    }
}

# ==============================================================================
# SECTION 3: CUSTOM MODEL IMPLEMENTATION PLACEHOLDERS
#
# NOTE: The following classes are placeholders for the models that require
# custom implementation. The actual implementation would involve significant
# code based on their respective research papers.
# ==============================================================================

class HGSUM(torch.nn.Module):
    def __init__(self, config):
        super().__init__()
        # HGSUM Implementation Details:
        # 1. Base text encoder-decoder initialized from PRIMERA.
        # 2. A custom graph construction module to build heterogeneous graphs
        #    (word, sentence, document nodes) from input text clusters.
        # 3. A Graph Attention Network (GAT) encoder to process the graph.
        # 4. A graph pooling/compressor module.
        # 5. A custom forward pass that integrates text and graph representations.
        # 6. A dual loss function (cross-entropy for generation + graph similarity).
        # This is a highly complex, research-level implementation.
        raise NotImplementedError("HGSUM requires a custom implementation based on its paper.")

    def forward(self, *args, **kwargs):
        pass

class DCA(torch.nn.Module):
    def __init__(self, config):
        super().__init__()
        # DCA (Deep Communicating Agents) Implementation Details:
        # 1. An encoder composed of multiple "agents" (e.g., LSTMs).
        # 2. The input document cluster is split, with each agent receiving one document.
        # 3. A communication mechanism where agents exchange hidden states.
        # 4. A single decoder with a contextual attention mechanism that attends
        #    over the outputs of all communicating agents.
        # 5. This architecture is fundamentally different from standard transformers
        #    and must be built from scratch using PyTorch modules.
        raise NotImplementedError("DCA requires a custom, non-transformer implementation.")

    def forward(self, *args, **kwargs):
        pass

class Absformer(torch.nn.Module):
    def __init__(self, config):
        super().__init__()
        # Absformer Implementation Details:
        # This is a two-phase unsupervised model.
        # Phase 1: Pre-train an encoder (e.g., DistilBERT) on the corpus with a
        #          Masked Language Modeling (MLM) objective. This encoder is then
        #          used to generate embeddings for document clustering.
        # Phase 2: Train a separate decoder. The decoder is trained to reconstruct
        #          documents from their embeddings (generated by the frozen encoder
        #          from Phase 1). For summarization, the decoder is then fed the
        #          cluster centroid embeddings to generate a summary.
        # This requires two separate training loops and careful management of model components.
        raise NotImplementedError("Absformer requires a custom two-phase unsupervised implementation.")

    def forward(self, *args, **kwargs):
        pass


# ==============================================================================
# SECTION 4: DATA HANDLING & PREPROCESSING
# ==============================================================================

def load_data():
    """Loads the CNN/Daily Mail dataset from Hugging Face."""
    logging.info(f"Loading dataset: {DATASET_NAME}, version {DATASET_VERSION}")
    try:
        dataset = load_dataset(DATASET_NAME, DATASET_VERSION)
        logging.info("Dataset loaded successfully.")
        logging.info(f"Dataset splits: {list(dataset.keys())}")
        logging.info(f"Train examples: {len(dataset['train'])}")
        logging.info(f"Validation examples: {len(dataset['validation'])}")
        logging.info(f"Test examples: {len(dataset['test'])}")
        return dataset
    except Exception as e:
        logging.error(f"Failed to load dataset: {e}")
        return None

def create_mds_clusters(dataset_split, k=3):
    """
    Creates synthetic multi-document clusters using TF-IDF and Nearest Neighbors.
    This is necessary for evaluating MDS-native models on an SDS dataset.
    """
    logging.info(f"Creating MDS clusters for dataset split of size {len(dataset_split)} with k={k}")
    vectorizer = TfidfVectorizer(max_features=5000, stop_words='english')
    tfidf_matrix = vectorizer.fit_transform(dataset_split['article'])

    nn = NearestNeighbors(n_neighbors=k, metric='cosine', algorithm='brute')
    nn.fit(tfidf_matrix)
    distances, indices = nn.kneighbors(tfidf_matrix)

    clustered_articles =
    for i in range(len(dataset_split)):
        neighbor_indices = indices[i]
        # Concatenate the anchor article with its k-1 neighbors
        cluster_text = " <EOD> ".join([dataset_split[int(j)]['article'] for j in neighbor_indices])
        clustered_articles.append(cluster_text)

    return clustered_articles

def preprocess_function_factory(tokenizer, config):
    """
    Factory to create a preprocessing function tailored to the model's config.
    """
    prefix = config.get("prefix", "")
    max_input_length = config.get("max_input_length", 1024)
    max_target_length = 128 # Standard for CNN/DM

    # Load spaCy model for entity extraction if needed
    nlp = None
    if config['type'] == 'sds_entity':
        logging.info("Loading spaCy model for entity extraction.")
        try:
            nlp = spacy.load("en_core_web_trf")
        except OSError:
            logging.warning("spaCy model 'en_core_web_trf' not found. Please run 'python -m spacy download en_core_web_trf'.")
            nlp = spacy.load("en_core_web_sm") # Fallback

    def preprocess_function(examples):
        """Tokenizes articles and highlights for training."""
        articles = examples['article']

        # --- Handle special preprocessing based on model type ---
        if config['type'] == 'sds_entity' and nlp:
            augmented_articles =
            for doc in nlp.pipe(articles):
                entities = [ent.text for ent in doc.ents]
                unique_entities = sorted(list(set(entities)))
                entity_prefix = f"ENTITIES: [{', '.join(unique_entities)}] "
                augmented_articles.append(entity_prefix + doc.text)
            articles = augmented_articles

        inputs = [prefix + doc for doc in articles]
        model_inputs = tokenizer(inputs, max_length=max_input_length, truncation=True, padding="max_length")

        # Tokenize targets
        with tokenizer.as_target_tokenizer():
            labels = tokenizer(examples["highlights"], max_length=max_target_length, truncation=True, padding="max_length")

        model_inputs["labels"] = labels["input_ids"]

        # For LED/PRIMERA, create global_attention_mask
        if config['model_class'] == LEDForConditionalGeneration:
            global_attention_mask = np.zeros_like(model_inputs['input_ids'])
            # Set global attention on the first token (<s>)
            global_attention_mask[:, 0] = 1
            model_inputs['global_attention_mask'] = global_attention_mask

        return model_inputs

    return preprocess_function

# ==============================================================================
# SECTION 5: EVALUATION
# ==============================================================================

# Load ROUGE metric
rouge_metric = evaluate.load("rouge")

def compute_metrics(eval_pred):
    """Computes ROUGE scores for a batch of predictions."""
    predictions, labels = eval_pred
    # Decode generated summaries into text
    decoded_preds = tokenizer.batch_decode(predictions, skip_special_tokens=True)
    # Replace -100 in the labels as we can't decode them
    labels = np.where(labels!= -100, labels, tokenizer.pad_token_id)
    # Decode reference summaries into text
    decoded_labels = tokenizer.batch_decode(labels, skip_special_tokens=True)

    # ROUGE expects a newline after each sentence
    decoded_preds = ["\n".join(pred.strip().split()) for pred in decoded_preds]
    decoded_labels = ["\n".join(label.strip().split()) for label in decoded_labels]

    # Compute ROUGE scores
    result = rouge_metric.compute(predictions=decoded_preds, references=decoded_labels, use_stemmer=False)
    result = {key: value * 100 for key, value in result.items()}

    # Add mean generated length
    prediction_lens = [np.count_nonzero(pred!= tokenizer.pad_token_id) for pred in predictions]
    result["gen_len"] = np.mean(prediction_lens)

    return {k: round(v, 4) for k, v in result.items()}


# ==============================================================================
# SECTION 6: TRAINING & INFERENCE PIPELINE
# ==============================================================================

def run_benchmark_for_model(model_name, config, dataset, cli_args):
    """
    Main function to run the full benchmark (train, infer, evaluate) for a single model.
    """
    logging.info(f"--- Starting Benchmark for Model: {model_name.upper()} ---")
    start_time = time.time()

    # --- 1. Handle Custom/Placeholder Models ---
    if config['model_class'] in:
        logging.warning(f"Model '{model_name}' requires a custom implementation and cannot be run automatically.")
        logging.warning("Skipping this model. Please implement the corresponding class to run.")
        # Return dummy results
        return {
            "Model": model_name.upper(),
            "ROUGE-1": "N/A", "ROUGE-2": "N/A", "ROUGE-L": "N/A",
            "Train Time (h)": "N/A", "GPU Model": torch.cuda.get_device_name(DEVICE) if torch.cuda.is_available() else "CPU",
            "Notes": "Custom implementation required."
        }

    # --- 2. Load Tokenizer and Model ---
    logging.info(f"Loading tokenizer and model from checkpoint: {config['checkpoint']}")
    global tokenizer # Make tokenizer global for compute_metrics
    tokenizer = config['tokenizer_class'].from_pretrained(config['checkpoint'])
    model = config['model_class'].from_pretrained(config['checkpoint']).to(DEVICE)

    # --- 3. Preprocess Data ---
    logging.info("Preprocessing data...")
    tokenized_dataset = dataset.copy()

    # Handle MDS clustering if required
    if config['type'] in ['mds', 'mds_custom']:
        logging.info("Applying MDS clustering...")
        for split in ['train', 'validation', 'test']:
            clustered_articles = create_mds_clusters(dataset[split])
            tokenized_dataset[split] = tokenized_dataset[split].remove_columns('article').add_column('article', clustered_articles)

    preprocess_function = preprocess_function_factory(tokenizer, config)
    tokenized_dataset = tokenized_dataset.map(preprocess_function, batched=True, remove_columns=dataset["train"].column_names)
    logging.info("Data preprocessing complete.")

    # --- 4. Fine-Tuning ---
    data_collator = DataCollatorForSeq2Seq(tokenizer=tokenizer, model=model)

    output_dir = os.path.join(cli_args.output_dir, model_name)
    training_args = Seq2SeqTrainingArguments(
        output_dir=output_dir,
        evaluation_strategy="epoch",
        learning_rate=float(cli_args.lr),
        per_device_train_batch_size=int(cli_args.batch_size),
        per_device_eval_batch_size=int(cli_args.batch_size),
        weight_decay=0.01,
        save_total_limit=3,
        num_train_epochs=int(cli_args.epochs),
        predict_with_generate=True,
        fp16=torch.cuda.is_available(), # Use mixed precision if on GPU
        push_to_hub=False,
        load_best_model_at_end=True, # Important for saving the best checkpoint
        metric_for_best_model="rougeL",
        report_to="none", # Disable wandb/tensorboard logging for simplicity
    )

    trainer = Seq2SeqTrainer(
        model=model,
        args=training_args,
        train_dataset=tokenized_dataset["train"],
        eval_dataset=tokenized_dataset["validation"],
        tokenizer=tokenizer,
        data_collator=data_collator,
        compute_metrics=compute_metrics,
    )

    logging.info(f"Starting training for {model_name}...")
    train_result = trainer.train()
    training_time_seconds = train_result.metrics["train_runtime"]
    logging.info(f"Training complete for {model_name}. Time taken: {training_time_seconds:.2f} seconds.")

    # Save the best model
    best_model_path = os.path.join(output_dir, "best_model")
    trainer.save_model(best_model_path)
    logging.info(f"Best model saved to {best_model_path}")

    # --- 5. Inference on Test Set ---
    logging.info("Running inference on the test set...")
    test_results = trainer.predict(tokenized_dataset["test"])
    decoded_preds = tokenizer.batch_decode(test_results.predictions, skip_special_tokens=True)

    # --- 6. Save Generated Summaries to.docx ---
    os.makedirs("results", exist_ok=True)
    doc_output_path = f"results/cnn_dailymail_{model_name}_test_summaries.docx"
    doc = Document()
    doc.add_heading(f"Test Set Summaries: {model_name.upper()}", 0)
    for i, pred in enumerate(decoded_preds):
        ref_summary = dataset['test'][i]['highlights']
        doc.add_heading(f"ID: {dataset['test'][i]['id']}", level=2)
        doc.add_paragraph(f"Reference Summary:\n{ref_summary}")
        doc.add_paragraph(f"Generated Summary:\n{pred}")
        doc.add_paragraph("-" * 50)
    doc.save(doc_output_path)
    logging.info(f"Generated summaries saved to {doc_output_path}")

    # --- 7. Final Evaluation & Result Logging ---
    final_metrics = test_results.metrics
    logging.info(f"Final ROUGE scores for {model_name}: {final_metrics}")

    end_time = time.time()
    total_time_hours = (end_time - start_time) / 3600.0

    return {
        "Model": model_name.upper(),
        "ROUGE-1": final_metrics.get("test_rouge1", 0.0),
        "ROUGE-2": final_metrics.get("test_rouge2", 0.0),
        "ROUGE-L": final_metrics.get("test_rougeL", 0.0),
        "Train Time (h)": round(total_time_hours, 2),
        "GPU Model": torch.cuda.get_device_name(DEVICE) if torch.cuda.is_available() else "CPU",
        "Notes": f"fp16, BS={cli_args.batch_size}, LR={cli_args.lr}"
    }


# ==============================================================================
# SECTION 7: MAIN EXECUTION
# ==============================================================================

def main(cli_args):
    """
    Main function to orchestrate the entire benchmark process.
    """
    # Load the dataset once
    dataset = load_data()
    if dataset is None:
        return

    all_results =

    # Loop through each model configuration and run the benchmark
    for model_name, config in MODEL_CONFIGS.items():
        # Allow running a single model via CLI argument
        if cli_args.model_name and model_name!= cli_args.model_name:
            continue

        # Update config with CLI args where applicable
        if cli_args.model_ckpt:
            config['checkpoint'] = cli_args.model_ckpt

        try:
            result = run_benchmark_for_model(model_name, config, dataset, cli_args)
            all_results.append(result)
        except Exception as e:
            logging.error(f"An error occurred while processing model {model_name}: {e}")
            all_results.append({
                "Model": model_name.upper(), "ROUGE-1": "Error", "ROUGE-2": "Error", "ROUGE-L": "Error",
                "Train Time (h)": "Error", "GPU Model": "Error", "Notes": str(e)
            })

    # --- Save aggregated results to Excel ---
    if all_results:
        results_df = pd.DataFrame(all_results)
        excel_path = "results_summary.xlsx"
        results_df.to_excel(excel_path, index=False)
        logging.info(f"Benchmark complete. Aggregated results saved to {excel_path}")
        print("\n--- Benchmark Summary ---")
        print(results_df.to_string())
    else:
        logging.warning("No models were run. Exiting.")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Run Abstractive Summarization Benchmark.")
    parser.add_argument("--epochs", default=3, help="Number of training epochs.")
    parser.add_argument("--batch_size", default=4, help="Training and evaluation batch size.")
    parser.add_argument("--lr", default="2e-5", help="Learning rate.")
    parser.add_argument("--model_ckpt", default=None, help="Optional: Override model checkpoint from Hugging Face.")
    parser.add_argument("--output_dir", default="checkpoints", help="Directory to save model checkpoints.")
    parser.add_argument("--model_name", default=None, help="Optional: Run benchmark for a single specified model name (e.g., 'bart').")

    args = parser.parse_args()
    main(args)

