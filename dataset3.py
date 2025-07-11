#                                                                         Code for xsum dataset

# ==============================================================================
#
# COMPLETE BENCHMARK SCRIPT FOR MULTI-DOCUMENT SUMMARIZATION ON XSUM
#
# This single file contains all the code necessary to run the benchmark,
# from data download to model training and final report generation.
#
#
# ==============================================================================

# --- Part 0: Imports ---
import os
import json
import logging
import time
import torch
import numpy as np
import pandas as pd
import evaluate
import nltk
from datasets import load_dataset, load_from_disk, Dataset, DatasetDict
from transformers import (
    AutoTokenizer,
    AutoModelForSeq2SeqLM,
    DataCollatorForSeq2Seq,
    Seq2SeqTrainingArguments,
    Seq2SeqTrainer,
    set_seed,
    pipeline,
    PreTrainedModel,
    PretrainedConfig,
)
from transformers.modeling_outputs import Seq2SeqLMOutput
from sentence_transformers import SentenceTransformer
from sklearn.cluster import AgglomerativeClustering
from sklearn.metrics.pairwise import cosine_similarity
from docx import Document
from openpyxl import Workbook
from tqdm import tqdm

# --- Part 1: Project Setup and Dependencies ---
def show_project_setup():
    """Prints the required dependencies to be saved in requirements.txt."""
    requirements = """
# Core ML/NLP Libraries
torch==2.1.0
transformers==4.35.2
datasets==2.15.0
evaluate==0.4.1
accelerate==0.25.0

# Data Processing & Clustering
scikit-learn==1.3.2
sentence-transformers==2.2.2
nltk==3.8.1

# File I/O for Reporting
openpyxl==3.1.2
python-docx==1.1.0

# Utility
tqdm==4.66.1
pandas==2.1.3
numpy==1.26.2
"""
    print("=" * 80)
    print("Project Setup: Create a file named 'requirements.txt' with the following content:")
    print("=" * 80)
    print(requirements)
    print("=" * 80)
    print("Install these dependencies using: pip install -r requirements.txt")
    print("=" * 80)
    
    # Download NLTK data needed for ROUGE scoring
    try:
        nltk.data.find('tokenizers/punkt')
    except nltk.downloader.DownloadError:
        print("Downloading NLTK 'punkt' model...")
        nltk.download('punkt')
        print("Download complete.")

# --- Part 2: Data Preparation and Corpus Creation ---
def prepare_data_pipeline(config):
    """
    Handles the entire data preparation process, from download and loading
    to clustering and saving the final pseudo-MDS corpus for XSUM.
    """
    
    # --- CORE FUNCTIONS ---
    def load_and_save_raw_xsum(data_dir):
        """Loads the XSUM dataset from Hugging Face and saves it locally."""
        if os.path.exists(data_dir):
            logging.info(f"Raw XSUM dataset already found at '{data_dir}'. Loading from disk.")
            return load_from_disk(data_dir)
            
        logging.info("Downloading and loading XSUM dataset from Hugging Face...")
        try:
            # XSUM is readily available from the datasets library
            dataset = load_dataset("xsum", name="default")
            logging.info(f"Saving raw XSUM dataset to '{data_dir}' for future runs.")
            dataset.save_to_disk(data_dir)
            return dataset
        except Exception as e:
            logging.error(f"Failed to load or save XSUM dataset. Error: {e}")
            return None

    def generate_embeddings(dataset, model, text_column='document'):
        """Generates sentence embeddings for each article in the dataset."""
        logging.info(f"Generating embeddings using '{config['embedding_model']}'...")
        corpus = dataset[text_column]
        embeddings = model.encode(corpus, show_progress_bar=True, convert_to_tensor=True)
        return embeddings.cpu().numpy()

    def perform_clustering(embeddings):
        """Performs Hierarchical Agglomerative Clustering on the embeddings."""
        logging.info(f"Performing Hierarchical Agglomerative Clustering with distance threshold {config['distance_threshold']}...")
        clustering_model = AgglomerativeClustering(
            n_clusters=None,
            distance_threshold=config['distance_threshold'],
            metric='cosine',
            linkage='average'
        )
        clustering_model.fit(embeddings)
        return clustering_model.labels_

    def create_mds_instances(dataset, embeddings, cluster_labels, text_column='document', summary_column='summary'):
        """Creates multi-document summarization instances from clustered articles."""
        logging.info("Creating pseudo-MDS instances...")
        clusters = {}
        for i, label in enumerate(cluster_labels):
            if label not in clusters:
                clusters[label] = []
            clusters[label].append(i)
        
        mds_data = {'document': [], 'summary': []}
        
        for cluster_id, indices in tqdm(clusters.items(), desc="Processing clusters"):
            if len(indices) < config['num_docs_in_cluster']:
                continue
            
            cluster_embeddings = embeddings[indices]
            
            # Heuristic: Select the most central article's summary as the "gold" summary
            centroid = np.mean(cluster_embeddings, axis=0)
            similarities = cosine_similarity(cluster_embeddings, centroid.reshape(1, -1))
            central_doc_local_idx = np.argmax(similarities)
            central_doc_global_idx = indices[central_doc_local_idx]
            docs_to_combine_indices = indices[:config['num_docs_in_cluster']]
            
            # Concatenate texts with a separator token
            combined_text = " <EOD> ".join([dataset[i][text_column] for i in docs_to_combine_indices])
            gold_summary = dataset[central_doc_global_idx][summary_column]
            
            mds_data['document'].append(combined_text)
            mds_data['summary'].append(gold_summary)
            
        return Dataset.from_dict(mds_data)

    # --- MAIN EXECUTION for data prep ---
    if os.path.exists(config['mds_corpus_path']):
        logging.info(f"Processed MDS corpus already exists at '{config['mds_corpus_path']}'. Skipping data preparation.")
        return

    device = 'cuda' if torch.cuda.is_available() else 'cpu'
    logging.info(f"Using device: {device}")
    
    embedding_model = SentenceTransformer(config['embedding_model'], device=device)
    raw_dataset = load_and_save_raw_xsum(config['sds_corpus_path'])
    
    if raw_dataset is None:
        return

    mds_splits = {}
    # Note: XSUM has 'train', 'validation', 'test' splits.
    for split_name in raw_dataset.keys():
        logging.info(f"\n--- Processing split: {split_name} ---")
        current_split_data = raw_dataset[split_name]
        
        embeddings = generate_embeddings(current_split_data, embedding_model, text_column='document')
        cluster_labels = perform_clustering(embeddings)
        mds_dataset_split = create_mds_instances(current_split_data, embeddings, cluster_labels, text_column='document', summary_column='summary')
        
        mds_splits[split_name] = mds_dataset_split
        logging.info(f"Created {len(mds_dataset_split)} MDS instances for the '{split_name}' split.")
        
    final_mds_dataset = DatasetDict(mds_splits)
    
    logging.info(f"Saving processed MDS corpus to '{config['mds_corpus_path']}'...")
    final_mds_dataset.save_to_disk(config['mds_corpus_path'])
    logging.info("Data preparation complete.")


# --- Part 3: Custom Model Architectures ---
# This section remains unchanged as it defines model architectures, not data handling.
# A simplified DCA implementation and placeholders for other complex models.
class DCAConfig(PretrainedConfig):
    model_type = "dca"
    def __init__(self, vocab_size=50265, hidden_size=256, num_agents=3, encoder_layers=2, decoder_layers=1, dropout=0.1, **kwargs):
        super().__init__(**kwargs)
        self.vocab_size = vocab_size
        self.hidden_size = hidden_size
        self.num_agents = num_agents
        self.encoder_layers = encoder_layers
        self.decoder_layers = decoder_layers
        self.dropout = dropout

class DCAModel(PreTrainedModel):
    config_class = DCAConfig
    # This is a simplified, conceptual implementation. A real one is more complex.
    pass # Full implementation omitted for brevity, assuming it's the same as the original.

class TGMultiSumModel(PreTrainedModel): # Placeholder
    pass

class AbsformerModel(PreTrainedModel): # Placeholder
    pass


# --- Part 4: Main Training and Evaluation Script ---
def run_benchmark_pipeline(config):
    """
    Orchestrates the fine-tuning and evaluation for all specified models on XSUM.
    """
    set_seed(config['seed'])
    rouge_metric = evaluate.load("rouge")
    
    def preprocess_function(examples, tokenizer, model_config):
        """Tokenizes the document and summary fields for XSUM."""
        max_input_length = model_config.get("max_input_length", 1024)
        max_target_length = model_config.get("max_target_length", 64) # Adjusted for XSUM
        prefix = model_config.get("kwargs", {}).get("prefix", "")

        # **MODIFIED**: Changed "text" to "document" for XSUM
        inputs = [prefix + doc for doc in examples["document"]]
        model_inputs = tokenizer(inputs, max_length=max_input_length, truncation=True, padding="max_length")
        
        with tokenizer.as_target_tokenizer():
            labels = tokenizer(examples["summary"], max_length=max_target_length, truncation=True, padding="max_length")
        
        model_inputs["labels"] = labels["input_ids"]
        
        if model_config.get("kwargs", {}).get("use_global_attention", False):
            global_attention_mask = np.zeros_like(model_inputs['input_ids'])
            global_attention_mask[:, 0] = 1 # Global attention on the first token
            model_inputs['global_attention_mask'] = list(global_attention_mask)
            
        return model_inputs

    def compute_metrics(eval_pred, tokenizer):
        """Computes ROUGE scores for evaluation."""
        predictions, labels = eval_pred
        decoded_preds = tokenizer.batch_decode(predictions, skip_special_tokens=True)
        
        labels = np.where(labels != -100, labels, tokenizer.pad_token_id)
        decoded_labels = tokenizer.batch_decode(labels, skip_special_tokens=True)

        decoded_preds = ["\n".join(nltk.sent_tokenize(pred.strip())) for pred in decoded_preds]
        decoded_labels = ["\n".join(nltk.sent_tokenize(label.strip())) for label in decoded_labels]

        result = rouge_metric.compute(predictions=decoded_preds, references=decoded_labels, use_stemmer=True)
        result = {key: value * 100 for key, value in result.items()}

        prediction_lens = [np.count_nonzero(pred != tokenizer.pad_token_id) for pred in predictions]
        result["gen_len"] = np.mean(prediction_lens)
        
        return {k: round(v, 4) for k, v in result.items()}

    # --- Main Training Loop ---
    logging.info("Loading datasets for training...")
    sds_dataset = load_from_disk(config['sds_corpus_path'])
    mds_dataset = load_from_disk(config['mds_corpus_path'])

    for model_name, model_config in config['models_to_run'].items():
        logging.info(f"\n{'='*25} RUNNING BENCHMARK FOR: {model_name} {'='*25}")
        
        dataset = mds_dataset if model_config.get("is_mds", False) else sds_dataset
        
        checkpoint = model_config["checkpoint"]
        
        if "is_custom" in model_config.get("kwargs", {}) and model_config["kwargs"]["is_custom"]:
             logging.warning(f"Custom model '{model_name}' not fully implemented. Skipping training.")
             continue
        else:
            logging.info(f"Loading tokenizer and model from checkpoint: {checkpoint}")
            tokenizer = AutoTokenizer.from_pretrained(checkpoint)
            model = AutoModelForSeq2SeqLM.from_pretrained(checkpoint)

        logging.info("Tokenizing dataset...")
        tokenized_datasets = dataset.map(
            lambda x: preprocess_function(x, tokenizer, model_config),
            batched=True,
            remove_columns=dataset["train"].column_names
        )
        
        data_collator = DataCollatorForSeq2Seq(tokenizer, model=model)
        model_output_dir = os.path.join(config['output_dir'], model_name)
        
        batch_size = model_config.get("batch_size", config['training_args']['per_device_train_batch_size'])
        training_args_dict = config['training_args'].copy()
        training_args_dict['per_device_train_batch_size'] = batch_size
        training_args_dict['per_device_eval_batch_size'] = batch_size
        
        training_args = Seq2SeqTrainingArguments(
            output_dir=model_output_dir,
            **training_args_dict
        )

        trainer = Seq2SeqTrainer(
            model=model,
            args=training_args,
            train_dataset=tokenized_datasets["train"].select(range(config['num_train_samples'])),
            eval_dataset=tokenized_datasets["validation"].select(range(config['num_eval_samples'])),
            tokenizer=tokenizer,
            data_collator=data_collator,
            compute_metrics=lambda p: compute_metrics(p, tokenizer),
        )

        logging.info(f"Starting fine-tuning for {model_name}...")
        try:
            train_result = trainer.train()
            trainer.save_model() 

            metrics = train_result.metrics
            metrics["train_samples"] = config['num_train_samples']
            trainer.log_metrics("train", metrics)
            trainer.save_metrics("train", metrics)
            trainer.save_state()

            logging.info(f"Evaluating {model_name} on the test set...")
            test_results = trainer.predict(tokenized_datasets["test"].select(range(config['num_test_samples'])))

            test_metrics = test_results.metrics
            test_metrics_path = os.path.join(model_output_dir, "test_results.json")
            with open(test_metrics_path, "w") as f:
                json.dump(test_metrics, f, indent=4)
            trainer.log_metrics("test", test_metrics)

        except Exception as e:
            logging.error(f"An error occurred for {model_name}: {e}", exc_info=True)

        logging.info(f"--- Finished benchmark for {model_name} ---")
    
    logging.info("All models have been processed.")


# --- Part 5: Results Generation and Reporting ---
def generate_final_reports(config):
    """
    Loads fine-tuned models, generates summaries on XSUM test set, and creates final reports.
    """
    if not os.path.exists(config['report_output_dir']):
        os.makedirs(config['report_output_dir'])
        
    sds_test_dataset = load_from_disk(config['sds_corpus_path'])['test'].select(range(config['num_report_samples']))
    mds_test_dataset = load_from_disk(config['mds_corpus_path'])['test'].select(range(config['num_report_samples']))
    
    all_model_metrics = []
    model_folders = [d for d in os.listdir(config['output_dir']) if os.path.isdir(os.path.join(config['output_dir'], d))]

    for model_name in tqdm(model_folders, desc="Generating reports for models"):
        model_path = os.path.join(config['output_dir'], model_name)
        
        try:
            model = AutoModelForSeq2SeqLM.from_pretrained(model_path)
            tokenizer = AutoTokenizer.from_pretrained(model_path)
        except Exception as e:
            logging.warning(f"Could not load model '{model_name}' from {model_path}. Skipping. Error: {e}")
            continue

        model_info = config['models_to_run'].get(model_name, {})
        test_dataset = mds_test_dataset if model_info.get("is_mds", False) else sds_test_dataset
        
        summarizer = pipeline("summarization", model=model, tokenizer=tokenizer, device=0 if torch.cuda.is_available() else -1)

        generated_summaries = []
        for example in tqdm(test_dataset, desc=f"Summarizing with {model_name}", leave=False):
            # **MODIFIED**: Changed "text" to "document" for XSUM
            input_text = example['document']
            if model_info.get("kwargs", {}).get("prefix"):
                input_text = model_info["kwargs"]["prefix"] + input_text

            summary = summarizer(input_text, max_length=128, min_length=20, do_sample=False)[0]['summary_text']
            
            generated_summaries.append({
                "original_text": example['document'],
                "reference_summary": example['summary'],
                "generated_summary": summary
            })

        # Save summaries to .docx
        doc = Document()
        doc.add_heading(f'Generated Summaries for {model_name} on XSUM', 0)
        for i, summary_info in enumerate(generated_summaries):
            doc.add_heading(f'Sample {i+1}', level=2)
            doc.add_paragraph(f"Original Document:\n{summary_info['original_text'][:1000]}...")
            doc.add_paragraph(f"\nReference Summary:\n{summary_info['reference_summary']}")
            doc.add_paragraph(f"\nGenerated Summary:\n{summary_info['generated_summary']}")
            doc.add_paragraph("-" * 20)
        # **MODIFIED**: Changed filename to "xsum_"
        doc.save(os.path.join(config['report_output_dir'], f"xsum_{model_name}_test_summaries.docx"))
        
        # Load metrics from training
        metrics_path = os.path.join(model_path, "test_results.json")
        if os.path.exists(metrics_path) and model_name in config['report_data']:
            with open(metrics_path, 'r') as f:
                metrics = json.load(f)
            
            report_info = config['report_data'][model_name]
            all_model_metrics.append({
                "Model Name": model_name,
                "Dataset": f"XSUM-MDS (k={config['num_docs_in_cluster']})" if model_info.get("is_mds") else "XSUM (k=1)",
                "ROUGE-1": metrics.get('test_rouge1', 'N/A'),
                "ROUGE-2": metrics.get('test_rouge2', 'N/A'),
                "ROUGE-L": metrics.get('test_rougeLsum', 'N/A'),
                "Training Time (h)": report_info['train_time_h'],
                "GPU/TPU used": report_info['gpu'],
                "Observations/Notes": report_info['notes']
            })

    # Create final Excel report
    if all_model_metrics:
        df = pd.DataFrame(all_model_metrics)
        # **MODIFIED**: Changed filename
        excel_path = os.path.join(config['report_output_dir'], "xsum_results_summary.xlsx")
        df.to_excel(excel_path, index=False)
        logging.info(f"Excel report saved to {excel_path}")

    logging.info(f"All reports successfully generated in '{config['report_output_dir']}'")

# --- Part 6: Main Execution Block ---
def main():
    """Main function to run the entire benchmark for XSUM."""

    # --- Global Configuration for XSUM ---
    config = {
        "seed": 42,
        # **MODIFIED**: Paths updated for XSUM
        "sds_corpus_path": "./xsum_data_raw",
        "mds_corpus_path": "./xsum_mds_corpus",
        "output_dir": "./benchmark_results_xsum",
        "report_output_dir": "./final_reports_xsum",
        
        "embedding_model": 'all-mpnet-base-v2',
        "num_docs_in_cluster": 3, # How many docs to combine for an MDS instance
        "distance_threshold": 0.7, # Cosine distance for clustering
        
        "num_train_samples": 1500, # Use a subset for faster demo runs
        "num_eval_samples": 300,
        "num_test_samples": 300,
        "num_report_samples": 50,
        
        # **MODIFIED**: Models to run, including an XSUM-specific baseline
        "models_to_run": {
            "PEGASUS-XSUM": {"checkpoint": "google/pegasus-xsum", "is_mds": False, "batch_size": 4, "max_target_length": 64},
            "BART": {"checkpoint": "facebook/bart-large-cnn", "is_mds": False, "batch_size": 4, "max_target_length": 64},
            "T5-base": {"checkpoint": "t5-base", "is_mds": False, "batch_size": 8, "kwargs": {"prefix": "summarize: "}},
            "PRIMERA": {"checkpoint": "allenai/PRIMERA-arxiv", "is_mds": True, "batch_size": 1, "kwargs": {"use_global_attention": True}, "max_input_length": 4096},
            "LED": {"checkpoint": "allenai/led-large-16384", "is_mds": False, "batch_size": 1, "kwargs": {"use_global_attention": True}, "max_input_length": 4096},
            "LongT5": {"checkpoint": "google/long-t5-tglobal-base", "is_mds": False, "batch_size": 2, "max_input_length": 4096},
            "DCA": {"checkpoint": "dca-custom", "is_mds": True, "batch_size": 2, "kwargs": {"is_custom": True}},
        },
        
        "training_args": {
            "num_train_epochs": 1,
            "per_device_train_batch_size": 2,
            "per_device_eval_batch_size": 2,
            "warmup_steps": 100,
            "weight_decay": 0.01,
            "logging_dir": "./logs_xsum",
            "logging_steps": 50,
            "evaluation_strategy": "epoch",
            "save_strategy": "epoch",
            "save_total_limit": 1,
            "predict_with_generate": True,
            "fp16": torch.cuda.is_available(),
            "load_best_model_at_end": True,
            "metric_for_best_model": "rouge2", # ROUGE-2 is often a key metric for XSUM
        },
        
        # **MODIFIED**: Placeholder data for the final report.
        # User should replace these with actual results after running the benchmark.
        "report_data": {
            "PEGASUS-XSUM": {"train_time_h": 4.5, "gpu": "A100-80GB", "notes": "Excellent baseline, already fine-tuned on XSUM. Sets a high bar."},
            "BART": {"train_time_h": 6.0, "gpu": "A100-80GB", "notes": "Strong baseline, but less abstractive than PEGASUS-XSUM by default."},
            "T5-base": {"train_time_h": 3.0, "gpu": "A100-80GB", "notes": "Fast to train but may struggle with the high abstractiveness of XSUM."},
            "PRIMERA": {"train_time_h": 12.0, "gpu": "A100-80GB", "notes": "Top performer on clustered inputs. Pre-training helps MDS task."},
            "LED": {"train_time_h": 11.5, "gpu": "A100-80GB", "notes": "Long context ability may be less crucial for the shorter docs in XSUM clusters."},
            "LongT5": {"train_time_h": 5.5, "gpu": "A100-80GB", "notes": "Efficient long-context model, good balance."},
            "DCA": {"train_time_h": 10.0, "gpu": "A100-80GB", "notes": "Legacy architecture, likely to underperform modern transformers."},
        }
    }

    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    # --- Run the full pipeline ---

    # 1. Show setup instructions
    show_project_setup()

    # 2. Prepare the data (download, cluster, save)
    logging.info("\n--- Starting Data Preparation Pipeline for XSUM ---")
    prepare_data_pipeline(config)

    # 3. Run the benchmark (fine-tune and evaluate all models)
    logging.info("\n--- Starting Model Training and Evaluation Pipeline on XSUM ---")
    run_benchmark_pipeline(config)

    # 4. Generate final reports (.docx summaries and .xlsx results)
    logging.info("\n--- Starting Final Report Generation for XSUM ---")
    generate_final_reports(config)

    logging.info("\nBenchmark execution finished.")

if __name__ == "__main__":
    main()
